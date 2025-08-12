# ===========================================
# app.py — Gestion de classe (Flask / Postgres)
# ===========================================
# Points clés :
# - Réglages UI (animation) stockés en DB (table app_settings)
# - Injection des réglages dans les templates via g + context_processor
# - Exports DOCX/PDF robustes (plusieurs fallbacks)
# - Routes: classes, élèves, évaluations, dictées, rapports, config/export
# -------------------------------------------

# ——— Standard library ———
import os
import re
import io
import csv
import shutil
import unicodedata
import tempfile
import subprocess
import time
from datetime import datetime, date
from pathlib import Path
import html as htmlmod
import json

# ——— Flask & DB ———
from flask import (
    Flask, current_app, abort, render_template, request,
    redirect, url_for, flash, jsonify, g
)
import psycopg2
import psycopg2.extras
from werkzeug.utils import secure_filename

# ——— DOCX / PDF ———
from docx import Document
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4 as RL_A4
from reportlab.lib.units import mm as RL_mm


# ===========================================
#  FLASK APP
# ===========================================
app = Flask(__name__)
app.secret_key = "supersecretkey"  # à remplacer en prod


# ===========================================
#  EXPORTS : Chemins & utilitaires
# ===========================================
# Chemins réseau (peuvent être surchargés par ENV: DOCS_ROOT_PRIMARY/SECONDARY)
PRIMARY_ROOT   = r"\\serveur\Documents\Education Nationale"   # chez toi (UNC)
SECONDARY_ROOT = r"Z:\Documents\Education Nationale"          # à l’école (RaiDrive)
REUNIONS_DIRNAME = "19 - Réunions"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REFERENCE_DOCX = os.path.join(BASE_DIR, "static", "export", "reference.docx")

# Cache de la racine choisie (évite d'interroger le FS tout le temps)
_ACTIVE_ROOT = None
_LAST_CHECK  = 0
_CHECK_EVERY_SEC = 30  # re-test toutes les 30s


def pick_docs_root() -> Path:
    """
    Retourne la racine disponible :
     - d'abord ENV DOCS_ROOT_PRIMARY / DOCS_ROOT_SECONDARY
     - sinon constantes PRIMARY_ROOT / SECONDARY_ROOT
     - sinon fallback local ./exports_local
    """
    global _ACTIVE_ROOT, _LAST_CHECK
    now = time.time()
    if _ACTIVE_ROOT and (now - _LAST_CHECK) < _CHECK_EVERY_SEC:
        return _ACTIVE_ROOT

    candidates = [
        os.getenv("DOCS_ROOT_PRIMARY")   or PRIMARY_ROOT,
        os.getenv("DOCS_ROOT_SECONDARY") or SECONDARY_ROOT,
    ]
    for raw in candidates:
        if not raw:
            continue
        p = Path(raw)
        if p.exists():
            _ACTIVE_ROOT = p
            _LAST_CHECK = now
            print(f"[EXPORT] racine sélectionnée => {p}")
            return p

    # Fallback local si rien d’accessible
    fb = Path(current_app.root_path) / "exports_local"
    fb.mkdir(parents=True, exist_ok=True)
    _ACTIVE_ROOT = fb
    _LAST_CHECK = now
    print(f"[EXPORT] aucune racine réseau dispo, fallback => {fb}")
    return fb


def find_year_folder(docs_root: Path, annee_scolaire: str) -> Path:
    """
    Trouve un dossier contenant '(AAAA-AAAA)'.
    Priorité : suffixe exact, sinon présence n'importe où, sinon racine.
    """
    wanted = f"({annee_scolaire})"
    best = None
    for child in docs_root.iterdir():
        if child.is_dir():
            name = child.name
            if name.endswith(wanted):
                best = child if (best is None) else (child if child.stat().st_mtime > best.stat().st_mtime else best)
    if best:
        return best
    for child in docs_root.iterdir():
        if child.is_dir() and wanted in child.name:
            return child
    return docs_root


def ensure_reunions_dir(year_dir: Path) -> Path:
    """Crée (si besoin) et retourne le dossier '19 - Réunions'."""
    d = year_dir / REUNIONS_DIRNAME
    d.mkdir(parents=True, exist_ok=True)
    return d


def resolve_export_dir(annee_scolaire: str, type_label: str | None, eleve_dirname: str | None) -> Path:
    """
    Dossier final d’export, selon ta structure :
      <root>\(AAAA-AAAA)\19 - Réunions\<Eleve OU Type>
    """
    root = pick_docs_root()
    year_dir = find_year_folder(root, annee_scolaire)
    reunions = ensure_reunions_dir(year_dir)
    final = reunions / (eleve_dirname or (type_label or "Autres réunions"))
    final.mkdir(parents=True, exist_ok=True)
    return final


def _slug(s: str) -> str:
    """Slug léger (enlève accents/diacritiques, sécurise)."""
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.category(c).startswith("M"))
    s = s.replace("/", "-").replace("\\", "-").strip()
    return s


def ensure_export_dir_for_rapport(conn, rapport_id: int, allow_create_year: bool = True) -> str:
    """
    Dossier final d’export pour un rapport :
      <root>\(AAAA-AAAA)\19 - Réunions\<NOM Prénom|Sous-type|Type>
    """
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute("""
            SELECT r.id, r.classe_id, r.eleve_id, r.titre, r.heure_debut,
                   rt.libelle AS type_lib, rst.libelle AS sous_lib,
                   c.annee AS classe_annee,
                   e.nom AS eleve_nom, e.prenom AS eleve_prenom
            FROM rapports r
            JOIN rapport_types rt ON r.type_id = rt.id
            LEFT JOIN rapport_sous_types rst ON r.sous_type_id = rst.id
            LEFT JOIN classes c ON r.classe_id = c.id
            LEFT JOIN eleves e ON r.eleve_id = e.id
            WHERE r.id = %s
        """, (rapport_id,))
        r = cur.fetchone()
        if not r:
            raise ValueError("Rapport introuvable")

        root = pick_docs_root()

        # 1) Dossier d’année
        if r["classe_id"] and r.get("classe_annee"):
            year_dir = find_year_folder(root, r["classe_annee"])
        else:
            year_dir = root / "Sans classe"
            year_dir.mkdir(exist_ok=True)

        # 2) Dossier '19 - Réunions'
        reunions = ensure_reunions_dir(year_dir)

        # 3) Sous-dossier
        if r["eleve_id"]:
            nom = (r["eleve_nom"] or "").upper().strip()
            prenom = (r["eleve_prenom"] or "").strip()
            sub = f"{_slug(nom)} {_slug(prenom)}".strip() or "Eleve"
        else:
            sub = _slug(r["sous_lib"] or r["type_lib"] or "Réunion")

        final = reunions / sub
        final.mkdir(parents=True, exist_ok=True)
        return str(final)


# ===========================================
#  DB : Connexion
# ===========================================
def get_db_connection():
    """
    Connexion Postgres.
    ⚠️ Adapte user/mdp/host/port selon ton environnement.
    """
    return psycopg2.connect(
        dbname="gestion_classe",
        user="postgres",
        password="kr6bkhe",
        host="localhost",
        port="5432"
    )


# ===========================================
#  Réglages UI (stockés en DB)
# ===========================================
ALLOWED_ANIM = {"slide", "fade", "scale", "wipe", "clip"}
DEFAULT_UI   = {"anim_mode": "slide", "anim_duration": 520}


def get_ui_settings_from_db(conn):
    """Lit 'ui' dans app_settings, avec garde-fous."""
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute("SELECT value FROM app_settings WHERE key=%s", ('ui',))
        row = cur.fetchone()

    if not row or not row["value"]:
        return DEFAULT_UI.copy()

    val = {**DEFAULT_UI, **row["value"]}
    mode = val.get("anim_mode", "slide")
    try:
        dur = int(val.get("anim_duration", 520))
    except (TypeError, ValueError):
        dur = 520

    if mode not in ALLOWED_ANIM:
        mode = "slide"
    dur = max(200, min(dur, 2000))
    return {"anim_mode": mode, "anim_duration": dur}


def set_ui_settings_in_db(conn, data: dict):
    """
    Écrit 'ui' dans app_settings (upsert), avec garde-fous.
    data = {"anim_mode": "...", "anim_duration": 520}
    """
    mode = data.get("anim_mode", DEFAULT_UI["anim_mode"])
    try:
        dur = int(data.get("anim_duration", DEFAULT_UI["anim_duration"]))
    except (TypeError, ValueError):
        dur = DEFAULT_UI["anim_duration"]

    if mode not in ALLOWED_ANIM:
        mode = "slide"
    dur = max(200, min(dur, 2000))

    with conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO app_settings(key, value)
                VALUES (%s, %s::jsonb)
                ON CONFLICT (key) DO UPDATE
                SET value = EXCLUDED.value, updated_at = now()
                """,
                ('ui', json.dumps({"anim_mode": mode, "anim_duration": dur}))
            )
    return {"anim_mode": mode, "anim_duration": dur}


@app.before_request
def load_ui_settings():
    """
    Avant chaque requête, on charge les réglages UI en g
    (pour pouvoir les injecter ensuite côté Jinja).
    """
    try:
        conn = get_db_connection()
        g.ui_settings = get_ui_settings_from_db(conn)
    except Exception:
        g.ui_settings = DEFAULT_UI.copy()
    finally:
        try:
            conn.close()
        except Exception:
            pass


# --- UI settings: DB first, fallback JSON ---
def _ui_from_everywhere():
    ui = DEFAULT_UI.copy()

    # Lecture DB uniquement
    try:
        conn = get_db_connection()
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute("SELECT value FROM app_settings WHERE key=%s", ('ui',))
            row = cur.fetchone()
        conn.close()
        if row and row.get("value"):
            ui.update(row["value"])
    except Exception:
        pass

    # Garde-fous
    mode = ui.get("anim_mode", "slide")
    dur  = int(ui.get("anim_duration", 520))
    if mode not in ALLOWED_ANIM:
        mode = "slide"
    ui["anim_mode"] = mode
    ui["anim_duration"] = max(200, min(dur, 2000))
    return ui

SETTINGS_PATH = os.path.join(BASE_DIR, "settings.json")

def load_settings():
    if os.path.exists(SETTINGS_PATH):
        with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"ui": {"anim_mode": "slide", "anim_duration": 520}}

def save_settings(data: dict):
    data = dict(data or {})
    data.setdefault("ui", {})
    with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)






@app.context_processor
def inject_ui():
    """
    Injecte 'ui' dans tous les templates :
    - tente la BDD
    - fallback sur settings.json
    - fallback sur DEFAULT_UI
    """
    ui = None
    conn = None
    try:
        conn = get_db_connection()
        ui = get_ui_settings_from_db(conn)  # merge + garde-fous dedans
    except Exception as e:
        print("[UI] fallback file/default (DB KO):", e)
        # ⇩⇩ lire settings.json en direct si présent, sinon fallback vide
        try:
            settings_path = os.path.join(BASE_DIR, "settings.json")
            with open(settings_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            ui = data.get("ui", {}) if isinstance(data, dict) else {}
        except Exception:
            ui = {}

    finally:
        try:
            if conn: conn.close()
        except Exception:
            pass

    # merge final avec defaults pour être sûr
    final_ui = dict(DEFAULT_UI, **(ui or {}))
    return {"ui": final_ui}






# ===========================================
#  UTILS HTML / EXPORT
# ===========================================
def _has_headings(html: str) -> bool:
    return bool(re.search(r'<h[1-6][\s>]', html or '', flags=re.I))


def _plain_text_from_html(html_str: str) -> str:
    """Retire le HTML et garde des retours simples (fallback texte)."""
    if not html_str:
        return ""
    txt = re.sub(r"</(p|div|h[1-6]|li|tr|br|table|ul|ol)>", "\n", html_str, flags=re.I)
    txt = re.sub(r"<[^>]+>", "", txt)
    txt = htmlmod.unescape(txt)
    txt = re.sub(r"\n\s*\n\s*\n+", "\n\n", txt).strip()
    return txt


def normalize_html_for_docx(html_in: str) -> str:
    """
    Normalise HTML “éditeur” pour aider les exports :
    - <p><strong>court</strong></p> → <h2>court</h2>
    - p[font-size] → h1/h2/p selon la taille
    - <br><br> → nouvelle ligne de paragraphe
    """
    if not html_in:
        return ""
    s = html_in

    s = re.sub(
        r'<p[^>]*>\s*<strong>([^<]{1,120})</strong>\s*</p>',
        r'<h2>\1</h2>',
        s, flags=re.I
    )

    def _size_to_hx(m):
        val = m.group(1)
        try:
            v = float(val)
            if 'px' in m.group(0).lower():
                v = v * 0.75  # px -> pt approx
        except Exception:
            v = 0
        tag = 'h1' if v >= 18 else ('h2' if v >= 14 else 'p')
        content = m.group(2)
        return f'<{tag}>{content}</{tag}>'

    s = re.sub(
        r'<p[^>]*style="[^"]*font-size:\s*(\d+(?:\.\d+)?)(?:pt|px)[^"]*"[^>]*>(.*?)</p>',
        _size_to_hx, s, flags=re.I | re.S
    )

    s = re.sub(r'(?:<br\s*/?>\s*){2,}', '</p><p>', s, flags=re.I)
    return s


def _wrap_html_for_pdf(body_html: str, title="Rapport") -> str:
    """
    Entoure le HTML d’un head minimal + CSS export (si présent).
    """
    css_path = os.path.join(current_app.root_path, "static", "export", "export.css")
    try:
        with open(css_path, "r", encoding="utf-8") as f:
            extra_css = f.read()
    except Exception:
        extra_css = """
        @page { size: A4; margin: 20mm; }
        body { font-family: Arial, "Liberation Sans", "DejaVu Sans", sans-serif; font-size:12pt; line-height:1.5; }
        table { border-collapse: collapse; width:100%; }
        table, th, td { border:1px solid #ddd; } th, td { padding:6px 8px; vertical-align:top; }
        h1,h2,h3 { margin:.4em 0; }
        """
    return f"""<!doctype html>
<html lang="fr"><head>
<meta charset="utf-8"><title>{title}</title>
<style>{extra_css}</style>
</head><body>
{body_html or ''}
</body></html>"""


def export_docx_best_effort(html_input: str, out_path: str, reference_docx: str | None = None):
    """
    HTML -> DOCX robuste (essais successifs) :
      1) Pandoc CLI
      2) pypandoc
      3) html2docx
      4) fallback texte
    + écrit un fichier debug .html à côté du .docx
    """
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    html_in = (html_input or "").strip()
    if not html_in:
        doc = Document()
        doc.add_paragraph("(rapport vide)")
        doc.save(out_path)
        return

    # normaliser pour aider Pandoc (titres/styles)
    html_in = normalize_html_for_docx(html_in)

    # debug
    try:
        debug_html = out_path[:-5] + "_DEBUG.html"
        with open(debug_html, "w", encoding="utf-8") as f:
            f.write(html_in)
        print(f"[DOCX] debug html -> {debug_html}")
    except Exception as e:
        print(f"[DOCX] debug html save KO: {e}")

    try:
        print(f"[DOCX] has headings? {_has_headings(html_in)}")
    except Exception:
        pass
    if reference_docx and os.path.exists(reference_docx):
        print(f"[DOCX] using reference: {reference_docx}")
    else:
        print("[DOCX] no reference.docx")

    # 1) Pandoc CLI direct
    try:
        pandoc_bin = os.getenv("PANDOC_BIN") or shutil.which("pandoc") or r"C:\Program Files\Pandoc\pandoc.exe"
        if pandoc_bin and os.path.exists(pandoc_bin):
            with tempfile.TemporaryDirectory() as td:
                html_file = os.path.join(td, "in.html")
                with open(html_file, "w", encoding="utf-8") as f:
                    f.write(html_in)
                cmd = [pandoc_bin, "-f", "html", "-t", "docx", html_file, "-o", out_path]
                if reference_docx and os.path.exists(reference_docx):
                    cmd += ["--reference-doc", reference_docx]
                print(f"[DOCX] run: {' '.join(cmd)}")
                subprocess.run(cmd, check=True)
                if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                    print("[DOCX] OK via pandoc CLI")
                    return
                else:
                    print("[DOCX] pandoc CLI a produit un fichier vide")
        else:
            print("[DOCX] pandoc non trouvé (PATH/PANDOC_BIN)")
    except subprocess.CalledProcessError as e:
        print(f"[DOCX] pandoc CLI KO (code={e.returncode})")
    except Exception as e:
        print(f"[DOCX] pandoc CLI exception: {e}")

    # 2) pypandoc
    try:
        import pypandoc
        extra = []
        if reference_docx and os.path.exists(reference_docx):
            extra += ["--reference-doc", reference_docx]
        pypandoc.convert_text(html_in, "docx", format="html", outputfile=out_path, extra_args=extra)
        if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            print("[DOCX] OK via pypandoc")
            return
        else:
            print("[DOCX] pypandoc a produit un fichier vide")
    except Exception as e:
        print(f"[DOCX] pypandoc KO: {e}")

    # 3) html2docx
    try:
        from html2docx import html2docx as _html2docx
        doc = Document()
        _html2docx(html_in, doc)
        all_text = "".join(p.text for p in doc.paragraphs).strip()
        if all_text or doc.tables:
            doc.save(out_path)
            print("[DOCX] OK via html2docx")
            return
        else:
            print("[DOCX] html2docx a produit “vide” -> fallback texte")
    except Exception as e:
        print(f"[DOCX] html2docx KO: {e}")

    # 4) Fallback texte
    text = _plain_text_from_html(html_in) or "(contenu non interprétable)"
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(out_path)
    print("[DOCX] fallback texte")


def export_pdf_faithful(html_body: str, out_path: str, title="Rapport"):
    """HTML -> PDF (Playwright, puis wkhtmltopdf, puis ReportLab fallback)."""
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    html_str = _wrap_html_for_pdf(html_body or "", title=title)

    # 1) Playwright/Chromium
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox"])
            page = browser.new_page()
            page.set_content(html_str, wait_until="networkidle")
            page.pdf(
                path=out_path,
                format="A4",
                margin={"top": "20mm", "right": "20mm", "bottom": "20mm", "left": "20mm"},
                print_background=True,
                prefer_css_page_size=True
            )
            browser.close()
            return
    except Exception as e:
        print(f"[PDF] Playwright KO: {e}")

    # 2) wkhtmltopdf via pdfkit
    try:
        import pdfkit
        wkhtml = (os.getenv("WKHTMLTOPDF_PATH") or shutil.which("wkhtmltopdf")
                  or r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
                  or r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe")
        cfg = pdfkit.configuration(wkhtmltopdf=wkhtml) if wkhtml and os.path.exists(wkhtml) else None
        options = {
            'quiet': '', 'encoding': 'UTF-8', 'page-size': 'A4',
            'margin-top': '20mm', 'margin-right': '20mm', 'margin-bottom': '20mm', 'margin-left': '20mm',
            'print-media-type': '', 'enable-local-file-access': '', 'dpi': 96, 'image-dpi': 300, 'image-quality': 92
        }
        pdfkit.from_string(html_str, out_path, configuration=cfg, options=options)
        return
    except Exception as e:
        print(f"[PDF] wkhtmltopdf KO: {e}")

    # 3) Fallback texte (ReportLab)
    try:
        text = _plain_text_from_html(html_body or "")
        c = rl_canvas.Canvas(out_path, pagesize=RL_A4)
        width, height = RL_A4
        x, y = 20 * RL_mm, height - 20 * RL_mm
        c.setFont("Helvetica", 11)
        for line in text.splitlines():
            if y < 20 * RL_mm:
                c.showPage(); c.setFont("Helvetica", 11); y = height - 20 * RL_mm
            c.drawString(x, y, line[:1200])
            y -= 14
        c.save()
    except Exception as e:
        print(f"[PDF] fallback KO: {e}")
        open(out_path, "wb").close()  # dernier recours


# ===========================================
#  HELPERS GÉNÉRAUX (réponses HTTP)
# ===========================================
@app.route('/static/style.css')
def style_css():
    """Permet d'utiliser un template Jinja pour générer du CSS."""
    return render_template('style.css.j2'), 200, {'Content-Type': 'text/css'}


@app.after_request
def add_header(response):
    """Empêche le cache agressif du navigateur sur les pages dynamiques."""
    response.cache_control.no_store = True
    return response


def parse_dt(s: str):
    """Parse tolérant sur quelques formats datetime/date usuels."""
    if not s:
        return None
    for fmt in ("%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            pass
    try:
        return datetime.fromisoformat(s)
    except Exception:
        return None


# ===========================================
#  PAGES
# ===========================================
@app.route("/")
def index():
    """Accueil : liste des classes."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    classes = cur.fetchall()

    # Ajoute liste des niveaux pour chaque classe
    for c in classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s ORDER BY niveau", (c["id"],))
        c["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    cur.close()
    conn.close()
    return render_template("index.html", classes=classes)


@app.route("/classe/<int:classe_id>")
def page_classe(classe_id):
    """
    Détail classe — plusieurs modes :
      - ?mode=eleves (par défaut)
      - ?mode=liste_evaluations (+ filtres)
      - ?mode=saisie_resultats&evaluation_id=...
      - ?mode=ajouter_rapport
      - ?mode=ajouter_dictee
    """
    mode = request.args.get("mode", "eleves")
    evaluation_id = request.args.get("evaluation_id")
    filtre_niveau = request.args.get("filtre_niveau")
    filtre_matiere = request.args.get("filtre_matiere")
    filtre_sous_matiere = request.args.get("filtre_sous_matiere")
    niveaux_filtres = []
    groupes_dict = {}

    # Année scolaire courante
    now = datetime.now()
    year = now.year
    month = now.month
    annee_scolaire = f"{year}-{year + 1}" if month >= 8 else f"{year - 1}-{year}"

    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Niveaux de la classe
    cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (classe_id,))
    niveaux_classe = [row["niveau"] for row in cur.fetchall()]

    # Toutes les classes (menu gauche)
    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    toutes_les_classes = cur.fetchall()
    for cl in toutes_les_classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    # Classe courante
    cur.execute("SELECT * FROM classes WHERE id = %s", (classe_id,))
    classe = cur.fetchone()
    if not classe:
        cur.close(); conn.close()
        flash("Classe introuvable.")
        return redirect(url_for("index"))
    classe["niveaux"] = niveaux_classe

    eleves = []
    evaluations = []
    matieres = []
    sous_matieres = []
    avancements = {}
    evaluation = objectifs = resultats = None
    eleves_par_niveau = None  # pour le mode 'ajouter_dictee'

    # Élèves
    cur.execute("SELECT * FROM eleves WHERE classe_id = %s", (classe_id,))
    eleves = cur.fetchall()

    # Matières et sous-matières
    cur.execute("SELECT * FROM matieres ORDER BY nom")
    matieres = cur.fetchall()
    cur.execute("SELECT * FROM sous_matieres ORDER BY nom")
    sous_matieres = cur.fetchall()

    # Mode : liste des évaluations (+ filtres)
    if mode == "liste_evaluations":
        requete = """
            SELECT e.id, e.titre, e.date, m.nom AS matiere, sm.nom AS sous_matiere
            FROM evaluations e
            LEFT JOIN matieres m ON e.matiere_id = m.id
            LEFT JOIN sous_matieres sm ON e.sous_matiere_id = sm.id
            WHERE e.classe_id = %s
        """
        params = [classe_id]

        if filtre_matiere:
            requete += " AND m.nom = %s"; params.append(filtre_matiere)
        if filtre_sous_matiere:
            requete += " AND sm.nom = %s"; params.append(filtre_sous_matiere)
        if filtre_niveau:
            requete += """
                AND EXISTS (
                    SELECT 1 FROM evaluations_niveaux en
                    WHERE en.evaluation_id = e.id AND en.niveau = %s
                )
            """
            params.append(filtre_niveau)

        requete += " ORDER BY e.date DESC"
        cur.execute(requete, params)
        evaluations = cur.fetchall()

        # Niveaux concernés par éval
        for ev in evaluations:
            cur.execute("SELECT niveau FROM evaluations_niveaux WHERE evaluation_id = %s", (ev["id"],))
            ev["niveaux"] = [row["niveau"] for row in cur.fetchall()]

        # Filtres dynamiques (niveaux présents dans les évals listées)
        evaluation_ids = [e["id"] for e in evaluations]
        if evaluation_ids:
            cur.execute("""
                SELECT DISTINCT niveau FROM evaluations_niveaux
                WHERE evaluation_id = ANY(%s)
                ORDER BY niveau
            """, (evaluation_ids,))
            niveaux_filtres = [row["niveau"] for row in cur.fetchall()]

        # Avancement des saisies (les absents '---' comptent comme complets)
        for ev in evaluations:
            cur.execute("SELECT COUNT(*) FROM objectifs WHERE evaluation_id = %s", (ev["id"],))
            nb_objectifs = cur.fetchone()["count"]

            total_attendus = nb_objectifs * len(eleves)
            if total_attendus == 0:
                avancements[ev["id"]] = 100
                continue

            cur.execute("SELECT * FROM resultats WHERE evaluation_id = %s", (ev["id"],))
            lignes = cur.fetchall()

            nb_complets = 0
            par_eleve = {}
            for r in lignes:
                par_eleve.setdefault(r["eleve_id"], {})[r["objectif_id"]] = r["niveau"]

            for el in eleves:
                reponses = par_eleve.get(el["id"], {})
                if len(reponses) == nb_objectifs and all(niv in ["NA", "PA", "A", "---"] for niv in reponses.values()):
                    nb_complets += 1

            progression = int((nb_complets / len(eleves)) * 100) if eleves else 0
            avancements[ev["id"]] = progression

    # Mode : saisie des résultats
    if mode == "saisie_resultats" and evaluation_id:
        cur.execute("SELECT * FROM evaluations WHERE id = %s", (evaluation_id,))
        evaluation = cur.fetchone()

        cur.execute("SELECT * FROM objectifs WHERE evaluation_id = %s", (evaluation_id,))
        objectifs = cur.fetchall()

        cur.execute("SELECT * FROM resultats WHERE evaluation_id = %s", (evaluation_id,))
        lignes = cur.fetchall()

        resultats = {eleve['id']: {} for eleve in eleves}
        for ligne in lignes:
            resultats[ligne["eleve_id"]][ligne["objectif_id"]] = ligne["niveau"]

    # Mode : ajouter rapport (précharge types / sous-types / élèves)
    if mode == "ajouter_rapport":
        cur.execute("SELECT id, code, libelle FROM rapport_types ORDER BY libelle;")
        types = cur.fetchall()

        default_type_id = None
        for t in types:
            if t["code"] == "entretien_parents":
                default_type_id = t["id"]; break
        if default_type_id is None and types:
            default_type_id = types[0]["id"]

        sous_types = []
        if default_type_id:
            cur.execute("""
                SELECT id, code, libelle
                FROM rapport_sous_types
                WHERE type_id = %s
                ORDER BY libelle
            """, (default_type_id,))
            sous_types = cur.fetchall()

        cur.execute("""
            SELECT id, prenom, nom
            FROM eleves
            WHERE classe_id = %s
            ORDER BY prenom, nom
        """, (classe_id,))
        eleves_classe = cur.fetchall()
    else:
        types = []; sous_types = []; eleves_classe = []

    # Mode : ajouter dictée (prépare structures)
    if mode == "ajouter_dictee":
        from collections import defaultdict
        eleves_par_niveau = {}
        ids_eleves = []

        for niveau in niveaux_classe:
            cur.execute("""
                SELECT * FROM eleves
                WHERE classe_id = %s AND niveau = %s
                ORDER BY nom
            """, (classe_id, niveau))
            eleves_niveau = cur.fetchall()
            eleves_par_niveau[niveau] = eleves_niveau
            ids_eleves.extend([e['id'] for e in eleves_niveau])

        cur.execute("""
            SELECT d.*
            FROM dictees d
            JOIN classes_niveaux cn ON cn.id = d.niveau_id
            WHERE cn.classe_id = %s
            ORDER BY d.date
        """, (classe_id,))
        dictees = cur.fetchall()
        dictees_bilan = [d for d in dictees if d.get('type') == 'bilan']

        # changements de groupe
        if ids_eleves:
            cur.execute("""
                SELECT eleve_id, groupe, date_changement
                FROM groupes_eleves
                WHERE eleve_id = ANY(%s)
                ORDER BY eleve_id, date_changement
            """, (ids_eleves,))
            changements_groupe = cur.fetchall()
        else:
            changements_groupe = []

        groupes_par_eleve = defaultdict(list)
        for row in changements_groupe:
            groupes_par_eleve[row['eleve_id']].append({
                'groupe': row['groupe'],
                'date': row['date_changement']
            })

        # groupe retenu à la date de chaque dictée bilan
        from collections import defaultdict as ddict
        groupes_dict = ddict(dict)
        for dictee in dictees_bilan:
            date_dictee = dictee['date']
            if isinstance(date_dictee, datetime):
                date_dictee = date_dictee.date()

            for eleve_id, changements in groupes_par_eleve.items():
                groupe = 'G3'
                for ch in changements:
                    ch_dt = ch['date']
                    if isinstance(ch_dt, datetime):
                        ch_dt = ch_dt.date()
                    if ch_dt <= date_dictee:
                        groupe = ch['groupe']
                    else:
                        break
                groupes_dict[eleve_id][date_dictee.isoformat()] = groupe

    cur.close()
    conn.close()

    return render_template(
        "classe.html",
        classe=classe,
        eleves=eleves,
        toutes_les_classes=toutes_les_classes,
        matieres=matieres,
        sous_matieres=sous_matieres,
        evaluations=evaluations,
        evaluation=evaluation,
        objectifs=objectifs,
        resultats=resultats,
        avancements=avancements,
        annee_scolaire=annee_scolaire,
        mode=mode,
        niveaux_filtres=niveaux_filtres,
        eleves_par_niveau=eleves_par_niveau,
        groupes_dict=groupes_dict,
        types=types,
        sous_types=sous_types,
        eleves_classe=eleves_classe
    )


# ===========================================
#  API DICTÉES
# ===========================================
@app.post("/api/dictees")
def api_save_dictee():
    """
    Upsert d’une dictée + (optionnel) ses résultats.
    Payload minimal :
      { classe_id, niveau, date, type("simple"|"bilan"), ... }
    """
    data = request.get_json(silent=True) or {}
    conn = None; cur = None
    try:
        classe_id = int(data.get("classe_id"))
        niveau_txt = data.get("niveau")
        ddate = data.get("date")                 # "YYYY-MM-DD" ou "YYYY-MM-DDTHH:MM"
        dtype = data.get("type")                 # "simple" | "bilan"

        # Normalisation date -> datetime
        if isinstance(ddate, str):
            ddate = ddate.strip()
            try:
                if len(ddate) == 10:  # YYYY-MM-DD
                    ddate = datetime.fromisoformat(ddate + "T00:00:00")
                else:
                    ddate = datetime.fromisoformat(ddate.replace(" ", "T"))
            except ValueError:
                try:
                    ddate = datetime.strptime(ddate, "%Y-%m-%d %H:%M:%S")
                except ValueError:
                    return jsonify(ok=False, error=f"Format de date invalide: {data.get('date')}"), 400
        elif not isinstance(ddate, datetime):
            return jsonify(ok=False, error="Champ 'date' manquant ou invalide"), 400

        if not (classe_id and niveau_txt and ddate and dtype in ("simple", "bilan")):
            return jsonify(ok=False, error="Payload incomplet ou invalide"), 400

        # savoir quelles clés sont présentes (pour ne pas écraser lors d’un lock-only)
        keys_present = set(data.keys())
        has_simple = "nb_mots_simple" in keys_present
        has_g1     = "nb_mots_g1"     in keys_present
        has_g2     = "nb_mots_g2"     in keys_present
        has_g3     = "nb_mots_g3"     in keys_present

        nb_simple = data.get("nb_mots_simple")
        nb_g1 = data.get("nb_mots_g1")
        nb_g2 = data.get("nb_mots_g2")
        nb_g3 = data.get("nb_mots_g3")

        resultats = data.get("resultats") or []
        if not isinstance(resultats, list):
            return jsonify(ok=False, error="`resultats` doit être une liste"), 400

        verrouille_payload = data.get("verrouille")

        conn = get_db_connection(); cur = conn.cursor()

        # 1) Résoudre niveau_id
        cur.execute("""
            SELECT cn.id
            FROM classes_niveaux cn
            WHERE cn.classe_id = %s AND cn.niveau = %s
            LIMIT 1
        """, (classe_id, niveau_txt))
        row = cur.fetchone()
        if not row:
            conn.rollback(); cur.close(); conn.close()
            return jsonify(ok=False, error=f"Niveau introuvable pour classe_id={classe_id}, niveau='{niveau_txt}'"), 404
        niveau_id = row[0]

        # 1bis) verrouille final
        dictee_id_payload = data.get("dictee_id")
        if dictee_id_payload:
            cur.execute("SELECT verrouille FROM dictees WHERE id = %s", (dictee_id_payload,))
            existing = cur.fetchone()
            existing_lock = existing[0] if existing else False
            verrouille_final = bool(verrouille_payload) if verrouille_payload is not None else existing_lock
        else:
            verrouille_final = bool(verrouille_payload) if verrouille_payload is not None else False

        # 2) Insert/Update dictée
        if dictee_id_payload:
            # UPDATE : conserver valeurs existantes si absentes du payload
            cur.execute("""
                SELECT nb_mots_simple, nb_mots_g1, nb_mots_g2, nb_mots_g3
                FROM dictees WHERE id = %s
            """, (dictee_id_payload,))
            row_existing = cur.fetchone()
            ex_simple, ex_g1, ex_g2, ex_g3 = row_existing if row_existing else (None, None, None, None)

            nb_simple_final = nb_simple if has_simple else ex_simple
            nb_g1_final     = nb_g1     if has_g1     else ex_g1
            nb_g2_final     = nb_g2     if has_g2     else ex_g2
            nb_g3_final     = nb_g3     if has_g3     else ex_g3

            sql = """
            UPDATE dictees
            SET niveau_id      = %s,
                date           = %s::date + (date::time),  -- change le JOUR, garde l'HEURE
                type           = %s,
                nb_mots_simple = %s,
                nb_mots_g1     = %s,
                nb_mots_g2     = %s,
                nb_mots_g3     = %s,
                verrouille     = %s,
                classe_id      = %s
            WHERE id = %s
            RETURNING id
            """
            params = (
                niveau_id, ddate, dtype,
                nb_simple_final, nb_g1_final, nb_g2_final, nb_g3_final,
                verrouille_final, classe_id, dictee_id_payload
            )
            cur.execute(sql, params)
            row = cur.fetchone()
            if not row:
                conn.rollback(); cur.close(); conn.close()
                return jsonify(ok=False, error=f"Dictée id={dictee_id_payload} introuvable pour update"), 404
            dictee_id = row[0]
            mode = "update"
        else:
            # INSERT
            date_dt = data.get("date_dt")  # peut être None
            sql = """
            INSERT INTO dictees (
              niveau_id, date, type,
              nb_mots_simple, nb_mots_g1, nb_mots_g2, nb_mots_g3,
              verrouille, classe_id
            ) VALUES (
              %s,
              COALESCE(
                %s::timestamptz AT TIME ZONE 'Europe/Paris',
                %s::date + (NOW() AT TIME ZONE 'Europe/Paris')::time
              ),
              %s, %s, %s, %s, %s, %s, %s
            )
            RETURNING id
            """
            params = (
                niveau_id, date_dt, ddate, dtype,
                nb_simple, nb_g1, nb_g2, nb_g3,
                verrouille_final, classe_id
            )
            cur.execute(sql, params)
            dictee_id = cur.fetchone()[0]
            mode = "insert"

        # 3) Upsert des résultats (si fournis)
        insert_sql = """
            INSERT INTO dictee_resultats (dictee_id, eleve_id, groupe, erreurs, nb_mots)
            VALUES (%s,%s,%s,%s,%s)
            ON CONFLICT (dictee_id, eleve_id) DO UPDATE
            SET groupe  = EXCLUDED.groupe,
                erreurs = EXCLUDED.erreurs,
                nb_mots = EXCLUDED.nb_mots
        """
        for r in resultats:
            eleve_id = r.get("eleve_id")
            if eleve_id is None:
                continue
            eleve_id = int(eleve_id)

            raw_err = r.get("erreurs")
            if raw_err in ("", None):
                erreurs = None
            else:
                try:
                    erreurs = int(raw_err)
                    if erreurs < 0:
                        erreurs = None
                except (TypeError, ValueError):
                    erreurs = None

            try:
                nb_mots_res = int(r.get("nb_mots") or 0)
                if nb_mots_res < 0:
                    nb_mots_res = 0
            except (TypeError, ValueError):
                nb_mots_res = 0

            groupe = r.get("groupe")  # "G1"/"G2"/"G3" ou None
            cur.execute(insert_sql, (dictee_id, eleve_id, groupe, erreurs, nb_mots_res))

        conn.commit()
        print(f"[api/dictees] COMMIT id={dictee_id} mode={mode}")
        cur.close(); conn.close()
        return jsonify(ok=True, dictee_id=dictee_id, mode=mode)

    except Exception as e:
        try:
            if conn: conn.rollback()
            if cur: cur.close()
            if conn: conn.close()
        except Exception:
            pass
        print(f"[api/dictees] ERROR: {e}")
        return jsonify(ok=False, error=str(e)), 500


@app.get("/api/dictees")
def api_list_dictees():
    """Liste des dictées + résultats, pour une classe donnée."""
    classe_id = int(request.args["classe_id"])
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("""
        SELECT d.id, d.date, d.type,
               d.nb_mots_simple, d.nb_mots_g1, d.nb_mots_g2, d.nb_mots_g3,
               d.verrouille,
               cn.niveau
        FROM dictees d
        JOIN classes_niveaux cn ON cn.id = d.niveau_id
        WHERE cn.classe_id = %s
        ORDER BY d.date ASC, d.id ASC
    """, (classe_id,))
    rows = cur.fetchall()

    dictees = []
    for r in rows:
        dt = r["date"]
        if isinstance(dt, datetime):
            date_for_input = dt.strftime("%Y-%m-%dT%H:%M")
            date_display = dt.strftime("%Y-%m-%d %H:%M:%S")
        elif isinstance(dt, date):
            date_for_input = dt.strftime("%Y-%m-%d")
            date_display = dt.strftime("%Y-%m-%d 00:00:00")
        else:
            date_for_input = str(dt)[:16].replace(" ", "T")
            date_display = str(dt)

        dictees.append({
            "id": r["id"],
            "niveau": r["niveau"],
            "type": r["type"],
            "nb_mots_simple": r["nb_mots_simple"],
            "nb_mots_g1": r["nb_mots_g1"],
            "nb_mots_g2": r["nb_mots_g2"],
            "nb_mots_g3": r["nb_mots_g3"],
            "verrouille": r["verrouille"],
            "date": date_display,
            "date_dt": date_for_input,
        })

    ids = [d["id"] for d in dictees]
    resultats_map = {}
    if ids:
        cur.execute("""
            SELECT dictee_id, eleve_id, groupe, erreurs, nb_mots
            FROM dictee_resultats
            WHERE dictee_id = ANY(%s)
        """, (ids,))
        for r in cur.fetchall():
            resultats_map.setdefault(r["dictee_id"], {})[r["eleve_id"]] = r

    cur.close(); conn.close()
    return jsonify(ok=True, dictees=dictees, resultats=resultats_map)


# ===========================================
#  CLASSES / ÉLÈVES / ÉVALS
# ===========================================
@app.route("/ajouter_classe", methods=["POST"])
def ajouter_classe():
    """Crée une classe + ses niveaux (évite doublons si déjà existants)."""
    niveaux = request.form.getlist("niveau")
    annee_debut = request.form.get("annee_debut")

    if not niveaux or not annee_debut:
        flash("Merci de remplir tous les champs.", "warning")
        return redirect(request.referrer)

    annee = f"{annee_debut}-{int(annee_debut) + 1}"

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Si classe même année existe déjà ?
        cur.execute("SELECT id FROM classes WHERE annee = %s", (annee,))
        classe_existante = cur.fetchone()

        if classe_existante:
            classe_id = classe_existante[0]
            # Niveaux déjà tous présents ?
            cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (classe_id,))
            niveaux_existants = [row[0] for row in cur.fetchall()]
            if set(niveaux_existants) == set(niveaux):
                flash(f"La classe {', '.join(niveaux)} {annee} existe déjà.", "danger")
                return redirect(request.referrer)

        # Nouvelle classe
        cur.execute("INSERT INTO classes (annee) VALUES (%s) RETURNING id", (annee,))
        classe_id = cur.fetchone()[0]

        # Niveaux associés
        for niveau in niveaux:
            cur.execute("INSERT INTO classes_niveaux (classe_id, niveau) VALUES (%s, %s)", (classe_id, niveau))

        conn.commit()
        flash("Classe créée avec succès.", "success")
    except Exception as e:
        conn.rollback()
        flash(f"Erreur : {e}", "danger")
    finally:
        conn.close()

    return redirect(request.referrer)


@app.route("/ajouter_eleve/<int:classe_id>", methods=["POST"])
def ajouter_eleve(classe_id):
    """Ajoute un élève à la classe."""
    nom = request.form["nom"]
    prenom = request.form["prenom"]
    niveau = request.form.get("niveau", "")
    date_naissance = request.form.get("date_naissance", "")

    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("""
        INSERT INTO eleves (nom, prenom, niveau, date_naissance, classe_id)
        VALUES (%s, %s, %s, %s, %s)
    """, (nom, prenom, niveau, date_naissance, classe_id))
    conn.commit()
    cur.close(); conn.close()
    return redirect(url_for("page_classe", classe_id=classe_id))


@app.route("/classe/<int:classe_id>/ajouter_evaluation", methods=["POST"])
def ajouter_evaluation(classe_id):
    """Crée une évaluation + objectifs + niveaux concernés."""
    titre = request.form.get("titre")
    date_str = request.form.get("date")
    matiere_nom = request.form.get("matiere")
    sous_matiere_nom = request.form.get("sous_matiere")
    objectifs = [o.strip() for o in request.form.getlist("objectifs[]") if o.strip()]
    niveaux_concernes = request.form.getlist("niveaux_concernes")

    if not titre or not date_str or not matiere_nom or not objectifs or not niveaux_concernes:
        flash("⚠️ Tous les champs obligatoires doivent être remplis.")
        return redirect(url_for("page_classe", classe_id=classe_id, mode="ajout_evaluation"))

    conn = get_db_connection(); cur = conn.cursor()
    try:
        # Matière
        cur.execute("SELECT id FROM matieres WHERE nom = %s", (matiere_nom,))
        matiere = cur.fetchone()
        if matiere:
            matiere_id = matiere[0]
        else:
            cur.execute("INSERT INTO matieres (nom) VALUES (%s)", (matiere_nom,))
            cur.execute("SELECT LASTVAL()")
            matiere_id = cur.fetchone()[0]

        # Sous-matière
        sous_matiere_id = None
        if sous_matiere_nom:
            cur.execute("SELECT id FROM sous_matieres WHERE nom = %s AND matiere_id = %s", (sous_matiere_nom, matiere_id))
            sm = cur.fetchone()
            if sm:
                sous_matiere_id = sm[0]
            else:
                cur.execute("INSERT INTO sous_matieres (nom, matiere_id) VALUES (%s, %s)", (sous_matiere_nom, matiere_id))
                cur.execute("SELECT LASTVAL()")
                sous_matiere_id = cur.fetchone()[0]

        # Évaluation
        cur.execute("""
            INSERT INTO evaluations (titre, date, classe_id, matiere_id, sous_matiere_id)
            VALUES (%s, %s, %s, %s, %s)
        """, (titre, date_str, classe_id, matiere_id, sous_matiere_id))
        cur.execute("SELECT LASTVAL()")
        evaluation_id = cur.fetchone()[0]

        # Objectifs
        for objectif in objectifs:
            cur.execute("INSERT INTO objectifs (evaluation_id, texte) VALUES (%s, %s)", (evaluation_id, objectif))

        # Niveaux
        for niveau in niveaux_concernes:
            cur.execute("INSERT INTO evaluations_niveaux (evaluation_id, niveau) VALUES (%s, %s)", (evaluation_id, niveau))

        conn.commit()
        flash("✅ Évaluation ajoutée avec succès !")

    except Exception as e:
        conn.rollback()
        flash("❌ Erreur : " + str(e))
    finally:
        cur.close(); conn.close()

    return redirect(url_for("page_classe", classe_id=classe_id, mode="liste_evaluations"))


@app.route("/evaluation/<int:evaluation_id>/modifier", methods=["GET", "POST"])
def modifier_evaluation(evaluation_id):
    """Formulaire de modification d’une évaluation (et ses liens)."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Évaluation
    cur.execute("""
        SELECT e.*, m.nom AS matiere_nom, sm.nom AS sous_matiere_nom, c.id AS classe_id, c.annee
        FROM evaluations e
        LEFT JOIN matieres m ON e.matiere_id = m.id
        LEFT JOIN sous_matieres sm ON e.sous_matiere_id = sm.id
        LEFT JOIN classes c ON e.classe_id = c.id
        WHERE e.id = %s
    """, (evaluation_id,))
    evaluation = cur.fetchone()
    if not evaluation:
        cur.close(); conn.close()
        abort(404)

    # Objectifs liés
    cur.execute("SELECT * FROM objectifs WHERE evaluation_id = %s ORDER BY id", (evaluation_id,))
    objectifs = cur.fetchall()

    # Niveaux de la classe
    cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (evaluation["classe_id"],))
    niveaux = [row["niveau"] for row in cur.fetchall()]

    # Niveaux concernés
    cur.execute("SELECT niveau FROM evaluations_niveaux WHERE evaluation_id = %s", (evaluation_id,))
    evaluation_niveaux = [row["niveau"] for row in cur.fetchall()]

    # Matières et sous-matières
    cur.execute("SELECT * FROM matieres ORDER BY nom")
    matieres = cur.fetchall()
    cur.execute("SELECT * FROM sous_matieres ORDER BY nom")
    sous_matieres = cur.fetchall()

    # Toutes les classes (menu)
    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    toutes_les_classes = cur.fetchall()
    for cl in toutes_les_classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    # Classe liée (bandeau titre)
    cur.execute("SELECT * FROM classes WHERE id = %s", (evaluation["classe_id"],))
    classe = cur.fetchone()
    cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (classe["id"],))
    classe["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    if request.method == "POST":
        titre = request.form.get("titre")
        date_str = request.form.get("date")
        matiere_nom = request.form.get("matiere")
        sous_matiere_nom = request.form.get("sous_matiere")
        objectifs_form = [obj.strip() for obj in request.form.getlist("objectifs[]") if obj.strip() != ""]
        niveaux_concernes = request.form.getlist("niveaux_concernes")

        try:
            # Matière
            cur.execute("SELECT id FROM matieres WHERE nom = %s", (matiere_nom,))
            matiere = cur.fetchone()
            if matiere:
                matiere_id = matiere["id"]
            else:
                cur.execute("INSERT INTO matieres (nom) VALUES (%s) RETURNING id", (matiere_nom,))
                matiere_id = cur.fetchone()["id"]

            # Sous-matière
            sous_matiere_id = None
            if sous_matiere_nom:
                cur.execute("SELECT id FROM sous_matieres WHERE nom = %s AND matiere_id = %s", (sous_matiere_nom, matiere_id))
                sm = cur.fetchone()
                if sm:
                    sous_matiere_id = sm["id"]
                else:
                    cur.execute("INSERT INTO sous_matieres (nom, matiere_id) VALUES (%s, %s) RETURNING id", (sous_matiere_nom, matiere_id))
                    sous_matiere_id = cur.fetchone()["id"]

            # Maj évaluation
            cur.execute("""
                UPDATE evaluations
                SET titre = %s, date = %s, matiere_id = %s, sous_matiere_id = %s
                WHERE id = %s
            """, (titre, date_str, matiere_id, sous_matiere_id, evaluation_id))

            # Objectifs
            cur.execute("DELETE FROM objectifs WHERE evaluation_id = %s", (evaluation_id,))
            for obj in objectifs_form:
                cur.execute("INSERT INTO objectifs (evaluation_id, texte) VALUES (%s, %s)", (evaluation_id, obj))

            # Niveaux concernés
            cur.execute("DELETE FROM evaluations_niveaux WHERE evaluation_id = %s", (evaluation_id,))
            for niv in niveaux_concernes:
                cur.execute("INSERT INTO evaluations_niveaux (evaluation_id, niveau) VALUES (%s, %s)", (evaluation_id, niv))

            conn.commit()
            flash("✅ Évaluation modifiée avec succès !")
            return redirect(url_for("page_classe", classe_id=evaluation["classe_id"], mode="liste_evaluations"))

        except Exception as e:
            conn.rollback()
            flash(f"❌ Erreur lors de la modification : {e}")

    cur.close(); conn.close()
    return render_template(
        "modifier_evaluation.html",
        evaluation=evaluation,
        objectifs=objectifs,
        niveaux=niveaux,
        matieres=matieres,
        sous_matieres=sous_matieres,
        toutes_les_classes=toutes_les_classes,
        classe=classe,
        evaluation_niveaux=evaluation_niveaux
    )


@app.route('/evaluation/<int:evaluation_id>/supprimer', methods=['POST'])
def supprimer_evaluation(evaluation_id):
    """Supprime évaluation + dépendances (objectifs, niveaux, résultats)."""
    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute("DELETE FROM resultats WHERE evaluation_id = %s", (evaluation_id,))
        cur.execute("DELETE FROM objectifs WHERE evaluation_id = %s", (evaluation_id,))
        cur.execute("DELETE FROM evaluations_niveaux WHERE evaluation_id = %s", (evaluation_id,))
        cur.execute("DELETE FROM evaluations WHERE id = %s", (evaluation_id,))
        conn.commit()
        flash("Évaluation et résultats supprimés avec succès.")
    except Exception as e:
        conn.rollback()
        flash(f"Erreur lors de la suppression : {e}")
    finally:
        cur.close(); conn.close()

    classe_id = request.form.get('classe_id')
    if classe_id:
        return redirect(url_for('page_classe', classe_id=classe_id, mode='liste_evaluations'))
    return redirect(url_for('index'))


@app.route("/importer_eleve_csv/<int:classe_id>", methods=["POST"])
def importer_eleve_csv(classe_id):
    """Import CSV élèves (Windows-1252 ; séparateur ';'). Initialise groupe G3."""
    if "csv_file" not in request.files or request.files["csv_file"].filename == "":
        flash("Aucun fichier sélectionné.")
        return redirect(url_for("page_classe", classe_id=classe_id))

    try:
        stream = io.StringIO(request.files["csv_file"].stream.read().decode("windows-1252"), newline=None)
        reader = csv.DictReader(stream, delimiter=";")

        conn = get_db_connection(); cur = conn.cursor()

        # mapping CSV -> DB
        mapping = {
            "Nom élève": "nom",
            "Prénom élève": "prenom",
            "Niveau": "niveau",
            "Cycle": "cycle",
            "Regroupement": "regroupement",
            "Classe": "classe",
            "Date inscription": "date_inscription",
            "Nom d'usage": "nom_usage",
            "Deuxième prénom": "deuxieme_prenom",
            "Troisième prénom": "troisieme_prenom",
            "Date naissance": "date_naissance",
            "Commune naissance": "commune_naissance",
            "Dépt naissance": "dept_naissance",
            "Pays naissance": "pays_naissance",
            "Sexe": "sexe",
            "Adresse": "adresse",
            "CP": "cp",
            "Commune": "commune",
            "Pays": "pays",
            "Etat": "etat"
        }

        for row in reader:
            colonnes = []; valeurs = []
            for csv_col, db_col in mapping.items():
                val = (row.get(csv_col) or "").strip()
                colonnes.append(db_col)
                valeurs.append(val if val != "" else None)

            # Ajout classe_id
            colonnes.append("classe_id"); valeurs.append(classe_id)

            sql = f"""
                INSERT INTO eleves ({', '.join(colonnes)})
                VALUES ({', '.join(['%s'] * len(valeurs))})
                RETURNING id
            """
            cur.execute(sql, valeurs)
            eleve_id = cur.fetchone()[0]

            # Groupe G3 par défaut
            cur.execute("""
                INSERT INTO groupes_eleves (eleve_id, groupe, date_changement)
                VALUES (%s, %s, %s)
            """, (eleve_id, 'G3', datetime.now().date()))

        conn.commit()
        flash("✅ Importation réussie.")
    except Exception as e:
        flash(f"❌ Erreur lors de l'import : {e}")
        raise e
    finally:
        try:
            cur.close(); conn.close()
        except Exception:
            pass

    return redirect(url_for("page_classe", classe_id=classe_id))


ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}


def allowed_file(filename):
    """Vérifie l’extension d’un fichier uploadé."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/classe/<int:classe_id>/eleve/<int:eleve_id>/changer_photo', methods=['POST'])
def changer_photo(classe_id, eleve_id):
    """Upload/sauvegarde la photo de l’élève (static/photos)."""
    if 'photo' not in request.files:
        flash("Aucun fichier sélectionné.")
        return redirect(url_for('detail_eleve', eleve_id=eleve_id, classe_id=classe_id))

    file = request.files['photo']
    if file.filename == '':
        flash("Aucun fichier sélectionné.")
        return redirect(url_for('detail_eleve', eleve_id=eleve_id, classe_id=classe_id))

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        upload_folder = os.path.join(current_app.root_path, 'static', 'photos')
        os.makedirs(upload_folder, exist_ok=True)
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)

        conn = get_db_connection(); cur = conn.cursor()
        cur.execute("UPDATE eleves SET photo_filename = %s WHERE id = %s", (filename, eleve_id))
        conn.commit()
        cur.close(); conn.close()

        flash("Photo mise à jour avec succès.")
    else:
        flash("Type de fichier non autorisé. Seules les images sont acceptées.")

    return redirect(url_for('detail_eleve', eleve_id=eleve_id, classe_id=classe_id))


@app.route("/debug/evaluations")
def debug_evaluations():
    """Petit listing brut des évaluations (debug)."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM evaluations ORDER BY id DESC")
    rows = cur.fetchall()
    cur.close(); conn.close()
    return "<pre>" + "\n".join([f"{r['id']} | Classe {r['classe_id']} | {r['titre']} ({r['date']})" for r in rows]) + "</pre>"


@app.route("/classe/<int:classe_id>/evaluation/<int:evaluation_id>/resultats", methods=["GET", "POST"])
def saisir_resultats(classe_id, evaluation_id):
    """Saisie des résultats d’une évaluation + calcul moyennes/appreciations."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    points_map = {'NA': 0, 'PA': 2, 'A': 4, '---': 0}

    def convertir_note(moyenne_points):
        note_sur_20 = (moyenne_points / 4) * 20
        if note_sur_20 == 20:           appreciation = 'D'
        elif 16 < note_sur_20 < 20:     appreciation = 'A+'
        elif 13 < note_sur_20 <= 16:    appreciation = 'A'
        elif 12 <= note_sur_20 <= 13:   appreciation = 'PA+'
        elif 8 <= note_sur_20 < 12:     appreciation = 'PA'
        elif 6 <= note_sur_20 < 8:      appreciation = 'PA-'
        else:                            appreciation = 'NA'
        return note_sur_20, appreciation

    # Classe & Evaluation
    cur.execute("SELECT * FROM classes WHERE id = %s", (classe_id,))
    classe = cur.fetchone()
    cur.execute("SELECT * FROM evaluations WHERE id = %s", (evaluation_id,))
    evaluation = cur.fetchone()

    # Niveaux concernés
    cur.execute("SELECT niveau FROM evaluations_niveaux WHERE evaluation_id = %s", (evaluation_id,))
    niveaux_concernes = [row['niveau'] for row in cur.fetchall()]

    # Filtre niveau
    niveau_selectionne = request.args.get('niveau')

    # Élèves (filtrés si niveau sélectionné)
    if niveau_selectionne and niveau_selectionne in niveaux_concernes:
        cur.execute("SELECT * FROM eleves WHERE classe_id = %s AND niveau = %s ORDER BY nom, prenom", (classe_id, niveau_selectionne))
    else:
        if niveaux_concernes:
            cur.execute("SELECT * FROM eleves WHERE classe_id = %s AND niveau = ANY(%s) ORDER BY niveau, nom, prenom", (classe_id, niveaux_concernes))
        else:
            cur.execute("SELECT * FROM eleves WHERE classe_id = %s ORDER BY nom, prenom", (classe_id,))
    eleves = cur.fetchall()

    # Objectifs
    cur.execute("SELECT * FROM objectifs WHERE evaluation_id = %s ORDER BY id", (evaluation_id,))
    objectifs = cur.fetchall()

    # Résultats existants -> dict
    resultats = {}
    cur.execute("SELECT * FROM resultats WHERE evaluation_id = %s", (evaluation_id,))
    for r in cur.fetchall():
        eid, oid = r["eleve_id"], r["objectif_id"]
        resultats.setdefault(eid, {})[oid] = r["niveau"]

    # Moyennes
    moyennes = {}
    for eleve in eleves:
        total_points = 0; nb_obj = 0
        for obj in objectifs:
            valeur = resultats.get(eleve["id"], {}).get(obj["id"], "---")
            if valeur in points_map:
                total_points += points_map[valeur]; nb_obj += 1
        moyenne_points = total_points / nb_obj if nb_obj > 0 else 0
        moyennes[eleve["id"]] = convertir_note(moyenne_points)

    absents = []

    if request.method == "POST":
        try:
            # Absents
            absents = [eleve["id"] for eleve in eleves if f"absent_{eleve['id']}" in request.form]

            for eleve in eleves:
                for objectif in objectifs:
                    champ = f"resultat_{eleve['id']}_{objectif['id']}"
                    if eleve["id"] in absents:
                        niveau = '---'  # absent => complet '---'
                    else:
                        niveau = request.form.get(champ)
                        if niveau not in ["NA", "PA", "A"]:
                            niveau = None  # pas de saisie -> ne rien écrire

                    cur.execute("SELECT id FROM resultats WHERE eleve_id = %s AND objectif_id = %s",
                                (eleve["id"], objectif["id"]))
                    existant = cur.fetchone()

                    if niveau in ["NA", "PA", "A", "---"]:
                        if existant:
                            cur.execute("UPDATE resultats SET niveau = %s WHERE id = %s", (niveau, existant["id"]))
                        else:
                            cur.execute(
                                "INSERT INTO resultats (eleve_id, objectif_id, evaluation_id, niveau) VALUES (%s, %s, %s, %s)",
                                (eleve["id"], objectif["id"], evaluation_id, niveau)
                            )
                    else:
                        if existant:
                            cur.execute("DELETE FROM resultats WHERE id = %s", (existant["id"],))

            conn.commit()
            flash("Résultats enregistrés.", "success")

            # rechargement résultats & moyennes
            resultats = {}
            cur.execute("SELECT * FROM resultats WHERE evaluation_id = %s", (evaluation_id,))
            for r in cur.fetchall():
                eid, oid = r["eleve_id"], r["objectif_id"]
                resultats.setdefault(eid, {})[oid] = r["niveau"]

            moyennes = {}
            for eleve in eleves:
                total_points = 0; nb_obj = 0
                for obj in objectifs:
                    valeur = resultats.get(eleve["id"], {}).get(obj["id"], "---")
                    if valeur in points_map:
                        total_points += points_map[valeur]; nb_obj += 1
                moyenne_points = total_points / nb_obj if nb_obj > 0 else 0
                moyennes[eleve["id"]] = convertir_note(moyenne_points)

        except Exception as e:
            conn.rollback()
            flash(f"Erreur : {e}", "danger")

    # Menu latéral
    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    toutes_les_classes = cur.fetchall()
    for cl in toutes_les_classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    cur.close(); conn.close()
    return render_template(
        "classe.html",
        mode="saisie_resultats",
        classe=classe,
        evaluation=evaluation,
        eleves=eleves,
        objectifs=objectifs,
        resultats=resultats,
        absents=absents,
        toutes_les_classes=toutes_les_classes,
        niveaux_concernes=niveaux_concernes,
        niveau_selectionne=niveau_selectionne,
        moyennes=moyennes
    )


@app.route('/classe/<int:classe_id>/eleve/<int:eleve_id>')
def detail_eleve(classe_id, eleve_id):
    """Petite fiche élève (avec calcul d’âge)."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Élève
    cur.execute("SELECT * FROM eleves WHERE id = %s", (eleve_id,))
    eleve = cur.fetchone()

    # Âge
    age = None
    if eleve and eleve.get('date_naissance'):
        birthdate = eleve['date_naissance']
        if isinstance(birthdate, str):
            try:
                birthdate = datetime.strptime(birthdate, "%Y-%m-%d").date()
            except ValueError:
                birthdate = None
        if isinstance(birthdate, date):
            today = date.today()
            age = today.year - birthdate.year - ((today.month, today.day) < (birthdate.month, birthdate.day))

    # Classe (pour titre/menu)
    cur.execute("SELECT * FROM classes WHERE id = %s", (classe_id,))
    classe = cur.fetchone()

    # Menu latéral
    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    toutes_les_classes = cur.fetchall()
    for cl in toutes_les_classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]

    cur.close(); conn.close()
    return render_template(
        "partials/detail_eleve.html",
        eleve=eleve,
        classe=classe,
        toutes_les_classes=toutes_les_classes,
        mode="detail_eleve",
        age=age,
    )


# ===========================================
#  GROUPES DICTÉES
# ===========================================
@app.route('/changer_groupe', methods=['POST'])
def changer_groupe():
    """Insère un changement de groupe daté pour un élève."""
    data = request.get_json()
    eleve_id = data.get('eleve_id')
    groupe = data.get('groupe')
    date_changement = data.get('date_changement')

    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("""
        INSERT INTO groupes_eleves (eleve_id, groupe, date_changement)
        VALUES (%s, %s, %s)
    """, (eleve_id, groupe, date_changement))
    conn.commit()
    cur.close(); conn.close()
    return jsonify({'status': 'ok'})


@app.route("/classe/<int:classe_id>/ajouter_dictee", methods=["GET", "POST"])
def ajouter_dictee(classe_id):
    """Page d’ajout de dictée (prépare datas côté template)."""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Classe
    cur.execute("SELECT * FROM classes WHERE id = %s", (classe_id,))
    classe = cur.fetchone()

    # Niveaux
    cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (classe_id,))
    niveaux = [row["niveau"] for row in cur.fetchall()]
    classe["niveaux"] = niveaux

    # Élèves par niveau
    eleves_par_niveau = {}
    ids_eleves = []
    for niveau in niveaux:
        cur.execute("""
            SELECT * FROM eleves
            WHERE classe_id = %s AND niveau = %s
            ORDER BY nom
        """, (classe_id, niveau))
        eleves = cur.fetchall()
        eleves_par_niveau[niveau] = eleves
        ids_eleves.extend([e["id"] for e in eleves])

    # Menu latéral
    cur.execute("SELECT id, annee FROM classes")
    classes = cur.fetchall()
    toutes_les_classes = []
    for cl in classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]
        toutes_les_classes.append(cl)

    # Groupes par élève (tous changements datés)
    from collections import defaultdict
    groupes_dict = defaultdict(dict)

    if ids_eleves:
        cur.execute("""
            SELECT eleve_id, groupe, date_changement
            FROM groupes_eleves
            WHERE eleve_id = ANY(%s)
            ORDER BY eleve_id, date_changement
        """, (ids_eleves,))
        changements = cur.fetchall()

        groupes_par_eleve = defaultdict(list)
        for row in changements:
            groupes_par_eleve[row["eleve_id"]].append({
                "date": row["date_changement"],
                "groupe": row["groupe"]
            })

        for eleve_id, changements in groupes_par_eleve.items():
            for ch in changements:
                date_str = ch["date"].strftime("%Y-%m-%d %H:%M:%S")
                groupes_dict[eleve_id][date_str] = ch["groupe"]

    cur.close(); conn.close()

    return render_template(
        "classe.html",
        mode="ajouter_dictee",
        classe=classe,
        toutes_les_classes=toutes_les_classes,
        eleves_par_niveau=eleves_par_niveau,
        groupes_dict=groupes_dict
    )


# ===========================================
#  RAPPORTS (CRUD + EXPORT)
# ===========================================
def fetch_rapport_types(conn):
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute("SELECT id, code, libelle FROM rapport_types ORDER BY libelle;")
        return cur.fetchall()


def fetch_rapport_sous_types(conn, type_id):
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute("""
            SELECT id, code, libelle
            FROM rapport_sous_types
            WHERE type_id = %s
            ORDER BY libelle;
        """, (type_id,))
        return cur.fetchall()


@app.route("/rapports/nouveau", methods=["GET"])
def nouveau_rapport():
    """UI de création d’un rapport (pré-remplit types/sous-types)."""
    classe_id = request.args.get("classe_id", type=int)
    eleve_id  = request.args.get("eleve_id",  type=int)

    conn = get_db_connection()
    try:
        types_ = fetch_rapport_types(conn)
        default_type_id = None
        for t in types_:
            if t["code"] == "entretien_parents":
                default_type_id = t["id"]; break
        if default_type_id is None and types_:
            default_type_id = types_[0]["id"]
        sous_types = fetch_rapport_sous_types(conn, default_type_id) if default_type_id else []
    finally:
        conn.close()

    return render_template("ajouter_rapport.html",
                           types=types_, sous_types=sous_types,
                           classe_id=classe_id, eleve_id=eleve_id)


@app.route("/api/rapports", methods=["POST"])
def api_create_rapport():
    """Crée un rapport vide, retourne son id + heure_debut."""
    data = request.get_json(force=True) or {}
    classe_id = data.get("classe_id")
    eleve_id  = data.get("eleve_id")
    type_id   = data.get("type_id")
    sous_type_id = data.get("sous_type_id")
    titre     = data.get("titre")

    if not type_id:
        return jsonify(ok=False, error="type_id manquant"), 400

    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute("""
            INSERT INTO rapports (classe_id, eleve_id, type_id, sous_type_id, titre, contenu, heure_debut)
            VALUES (%s,%s,%s,%s,%s,'', NOW())
            RETURNING id, heure_debut;
        """, (classe_id, eleve_id, type_id, sous_type_id, titre))
        rid, hdeb = cur.fetchone()
        conn.commit()
        return jsonify(ok=True, id=rid, heure_debut=hdeb.isoformat()), 201
    except Exception as e:
        conn.rollback()
        return jsonify(ok=False, error=str(e)), 500
    finally:
        cur.close(); conn.close()


@app.route("/api/rapports/<int:rapport_id>", methods=["PATCH"])
def api_update_rapport(rapport_id):
    """Patch partiel d’un rapport ; met à jour heure_fin à chaque écriture."""
    data = request.get_json(force=True) or {}
    champs, vals = [], []

    for cle in ("titre", "contenu", "type_id", "sous_type_id", "classe_id", "eleve_id"):
        if cle in data:
            champs.append(f"{cle}=%s")
            vals.append(data[cle])

    champs.append("heure_fin=NOW()")

    if not champs:
        return jsonify(ok=True)

    sql = f"UPDATE rapports SET {', '.join(champs)} WHERE id=%s RETURNING heure_fin, updated_at;"
    vals.append(rapport_id)

    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute(sql, tuple(vals))
        row = cur.fetchone()
        if not row:
            conn.rollback()
            return jsonify(ok=False, error="rapport introuvable"), 404
        conn.commit()
        heure_fin, updated_at = row
        return jsonify(ok=True, heure_fin=heure_fin.isoformat(), updated_at=updated_at.isoformat())
    except Exception as e:
        conn.rollback()
        return jsonify(ok=False, error=str(e)), 500
    finally:
        cur.close(); conn.close()


@app.route("/api/rapport_sous_types/<int:type_id>", methods=["GET"])
def api_get_sous_types(type_id):
    """Liste les sous-types pour un type donné."""
    conn = get_db_connection()
    try:
        sts = fetch_rapport_sous_types(conn, type_id)
        return jsonify(sts)
    finally:
        conn.close()


@app.post("/api/rapports/<int:rapport_id>/export")
def api_export_rapport(rapport_id):
    """
    Export d’un rapport en DOCX ou PDF.
    Body JSON :
      { "format": "docx"|"pdf", "html": "<override optionnel>" }
    """
    data = request.get_json(silent=True) or {}
    fmt = (data.get("format") or "docx").lower()
    html_override = (data.get("html") or "").strip()

    # Charger rapport
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""
        SELECT r.titre, r.contenu, r.classe_id,
               rt.libelle AS type_lib, rst.libelle AS sous_lib, c.annee AS classe_annee
        FROM rapports r
        JOIN rapport_types rt ON r.type_id = rt.id
        LEFT JOIN rapport_sous_types rst ON r.sous_type_id = rst.id
        LEFT JOIN classes c ON r.classe_id = c.id
        WHERE r.id = %s
    """, (rapport_id,))
    r = cur.fetchone()
    cur.close(); conn.close()
    if not r:
        return jsonify(ok=False, error="Rapport introuvable"), 404

    titre = r["titre"] or (r["sous_lib"] or r["type_lib"] or "Rapport")
    contenu_html = html_override if html_override else (r["contenu"] or "")

    # Dossier de sortie (selon structure)
    conn = get_db_connection()
    try:
        export_dir = ensure_export_dir_for_rapport(conn, rapport_id, allow_create_year=False)
    finally:
        conn.close()

    # Export
    safe_titre = "".join(ch if ch.isalnum() or ch in " -_." else "_" for ch in titre).strip()[:80] or "rapport"
    if fmt == "docx":
        out_path = os.path.join(export_dir, f"{safe_titre}.docx")
        reference_docx = os.path.join(current_app.root_path, "static", "export", "reference.docx")
        if not os.path.exists(reference_docx):
            reference_docx = None
        export_docx_best_effort(contenu_html, out_path, reference_docx=reference_docx)

    elif fmt == "pdf":
        out_path = os.path.join(export_dir, f"{safe_titre}.pdf")
        export_pdf_faithful(contenu_html, out_path, title=titre)
    else:
        return jsonify(ok=False, error="Format inconnu"), 400

    try:
        size = os.path.getsize(out_path)
    except Exception:
        size = None
    return jsonify(ok=True, path=out_path, size=size)


# ===========================================
#  CONFIG EXPORT (chemins + UI/animation)
# ===========================================
@app.route("/config/export", methods=["GET", "POST"])
def config_export():
    """
    Page de configuration :
      - Chemins d’export (en mémoire process, surchargés par ENV)
      - Réglages UI (anim_mode, anim_duration) stockés en DB (app_settings)
    """
    global PRIMARY_ROOT, SECONDARY_ROOT, REUNIONS_DIRNAME, _ACTIVE_ROOT, _LAST_CHECK

    # Menu latéral (classes)
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM classes ORDER BY annee DESC")
    toutes_les_classes = cur.fetchall()
    for cl in toutes_les_classes:
        cur.execute("SELECT niveau FROM classes_niveaux WHERE classe_id = %s", (cl["id"],))
        cl["niveaux"] = [row["niveau"] for row in cur.fetchall()]
    classe = toutes_les_classes[0] if toutes_les_classes else None
    cur.close()
    # ⚠️ on garde 'conn' ouvert pour lire/écrire les réglages UI juste après

    # Valeurs courantes (ENV > globals)
    current_primary   = os.getenv("DOCS_ROOT_PRIMARY")   or PRIMARY_ROOT
    current_secondary = os.getenv("DOCS_ROOT_SECONDARY") or SECONDARY_ROOT
    current_reunions  = REUNIONS_DIRNAME

    # UI depuis DB
    ui_settings = get_ui_settings_from_db(conn)
    current_anim_mode     = ui_settings["anim_mode"]
    current_anim_duration = ui_settings["anim_duration"]

    if request.method == "POST":
        # — chemins existants —
        new_primary   = (request.form.get("primary_root") or "").strip() or current_primary
        new_secondary = (request.form.get("secondary_root") or "").strip() or current_secondary
        new_reunions  = (request.form.get("reunions_dirname") or "").strip() or current_reunions

        PRIMARY_ROOT     = new_primary
        SECONDARY_ROOT   = new_secondary
        REUNIONS_DIRNAME = new_reunions

        # Force re-détection de la racine dispo
        _ACTIVE_ROOT = None; _LAST_CHECK = 0

        # — UI (animation) —
        mode = (request.form.get("anim_mode") or current_anim_mode).strip()
        try:
            duration = int(request.form.get("anim_duration", current_anim_duration))
        except ValueError:
            duration = current_anim_duration

        set_ui_settings_in_db(conn, {"anim_mode": mode, "anim_duration": duration})

        flash("Paramètres enregistrés.", "success")
        conn.close()
        return redirect(url_for(".config_export"))

    # GET
    conn.close()
    return render_template(
        "partials/config_export.html",
        primary_root=current_primary,
        secondary_root=current_secondary,
        reunions_dirname=current_reunions,
        current_anim_mode=current_anim_mode,
        current_anim_duration=current_anim_duration,
        toutes_les_classes=toutes_les_classes,
        classe=classe
    )


@app.get("/api/config/test-paths")
def api_test_paths():
    """Test simple d’existence des chemins configurés."""
    p = request.args.get("primary")   or PRIMARY_ROOT
    s = request.args.get("secondary") or SECONDARY_ROOT
    res = {
        "primary":   {"path": p, "exists": os.path.exists(p)},
        "secondary": {"path": s, "exists": os.path.exists(s)},
    }
    return jsonify(ok=True, **res)


# ===========================================
#  DÉMARRAGE LOCAL
# ===========================================
if __name__ == "__main__":
    # Test DB au boot (facultatif)
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT version();")
        version = cur.fetchone()
        print("✅ Connexion PostgreSQL OK :", version)
        conn.close()
    except Exception as e:
        print("❌ Erreur de connexion à PostgreSQL :", e)

    app.run(debug=True)
